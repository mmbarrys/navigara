from flask import Flask, jsonify, request, send_file 
from flask_cors import CORS
import networkx as nx
import json
import os
import re 
import datetime 
import uuid 
from dotenv import load_dotenv
import logging
from logging.handlers import RotatingFileHandler
import google.generativeai as genai
import requests
import fitz  # PyMuPDF
import docx 
from werkzeug.utils import secure_filename
from googleapiclient.discovery import build 
import markdown 
from xhtml2pdf import pisa 
from io import BytesIO 
from flask_sqlalchemy import SQLAlchemy 

# --- Konfigurasi Awal ---
load_dotenv(); app = Flask(__name__)
NETLIFY_APP_URL = "https://navigara.netlify.app" # GANTI
NGROK_TUNNEL_URL = "https://nonerroneously-unvoluptuous-alena.ngrok-free.dev" # GANTI
cors_origins = ["http://localhost:5173", NETLIFY_APP_URL, NGROK_TUNNEL_URL ]
CORS(app, origins=cors_origins, supports_credentials=True) 
UPLOAD_FOLDER = 'uploads'; os.makedirs(UPLOAD_FOLDER, exist_ok=True); app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///assessment_log.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# --- Model Database Log ---
class AssessmentLog(db.Model):
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    timestamp = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    module = db.Column(db.String(50)); candidate_name = db.Column(db.String(200)); jabatan_or_program = db.Column(db.String(200))
    skor_potensi = db.Column(db.Integer, nullable=True); skor_kinerja = db.Column(db.Integer, nullable=True)
    recommendation = db.Column(db.String(100), nullable=True); sentiment = db.Column(db.String(100), nullable=True)
with app.app_context(): db.create_all()

# --- Konfigurasi Logging ---
log_formatter = logging.Formatter('%(asctime)s %(levelname)s:%(name)s:%(funcName)s(%(lineno)d): %(message)s')
log_file = 'navigara_backend.log'
file_handler = RotatingFileHandler(log_file, maxBytes=1024*1024*5, backupCount=2); file_handler.setFormatter(log_formatter); file_handler.setLevel(logging.INFO)
console_handler = logging.StreamHandler(); console_handler.setFormatter(log_formatter); console_handler.setLevel(logging.INFO)
logging.basicConfig(level=logging.INFO, handlers=[file_handler, console_handler])
app.logger.info("Server NAVIGARA (Python) v3.8 Enhanced SKP Prompt Dimulai...")

# --- Konfigurasi Klien API ---
GEMINI_KEY = os.getenv('GEMINI_API_KEY'); BYTEPLUS_KEY = os.getenv('BYTEPLUS_API_KEY'); GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY'); GOOGLE_CSE_ID = os.getenv('GOOGLE_CSE_ID')
GEMINI_MODEL_NAME = "gemini-2.5-flash"; BYTEPLUS_MODEL_NAME = "seed-1-6-250615"; BYTEPLUS_API_ENDPOINT = "https://ark.ap-southeast.bytepluses.com/api/v3/chat/completions"
gemini_model = None; google_search_service = None 
if GEMINI_KEY:
    try: genai.configure(api_key=GEMINI_KEY); gemini_model = genai.GenerativeModel(GEMINI_MODEL_NAME); app.logger.info(f"Gemini OK: {GEMINI_MODEL_NAME}")
    except Exception as e: app.logger.error(f"Gemini Fail: {str(e)}")
else: app.logger.warning("GEMINI_API_KEY missing.")
if GOOGLE_API_KEY and GOOGLE_CSE_ID:
    try: google_search_service = build("customsearch", "v1", developerKey=GOOGLE_API_KEY); app.logger.info("Google Search OK.")
    except Exception as e: app.logger.error(f"Google Search Fail: {str(e)}")
else: app.logger.warning("GOOGLE keys missing. OSINT dibatasi.")
if not BYTEPLUS_KEY: app.logger.warning("BYTEPLUS_API_KEY missing.")

# --- Helper: Ekstraksi Teks File ---
def extract_text_from_file(file_path):
    try:
        ext = file_path.lower().split('.')[-1]
        text_content = None 
        if ext == 'pdf':
            doc = fitz.open(file_path); text_content = "".join(page.get_text() for page in doc); doc.close()
        elif ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f: text_content = f.read()
        elif ext == 'docx':
            doc = docx.Document(file_path); text_content = "\n".join([para.text for para in doc.paragraphs])
        elif ext == 'doc': 
             app.logger.warning(".doc might not be fully supported.")
             try: doc = docx.Document(file_path); text_content = "\n".join([para.text for para in doc.paragraphs])
             except: text_content = None 
        if text_content: app.logger.info(f"Ekstraksi teks dari {os.path.basename(file_path)} berhasil.")
        return text_content
    except Exception as e: app.logger.error(f"Error fatal ekstrak file {file_path}: {str(e)}"); return None

# --- Helper: Fungsi Panggilan AI ---
def call_gemini_api(prompt):
    if not gemini_model: return "Error: Klien API Gemini tidak terkonfigurasi."
    app.logger.info(f"Calling Gemini API ({len(prompt)} chars)")
    try: response = gemini_model.generate_content(prompt); return response.text
    except Exception as e: app.logger.error(f"Error API Gemini: {str(e)}"); return f"Error API Gemini: {str(e)}"
def call_byteplus_api(prompt, system_prompt="Anda asisten AI."):
    if not BYTEPLUS_KEY: return "Error: Klien API Byteplus tidak terkonfigurasi."
    app.logger.info(f"Calling Byteplus API ({len(prompt)} chars)")
    headers = {"Authorization": f"Bearer {BYTEPLUS_KEY}", "Content-Type": "application/json"}
    payload = {"model": BYTEPLUS_MODEL_NAME, "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}]}
    try:
        response = requests.post(BYTEPLUS_API_ENDPOINT, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content']
    except requests.exceptions.ReadTimeout: return "Error: API Byteplus Timeout (60s)."
    except Exception as e: app.logger.error(f"Error API Byteplus: {str(e)}"); return f"Error API Byteplus: {str(e)}"

# --- Helper: OSINT Engine ---
def run_osint_analysis(query):
    app.logger.info(f"OSINT Engine: Google Search '{query}'")
    articles = []
    if not google_search_service:
        articles.append({"source": "Sistem", "title": "Google Search API Error", "url": "#", "snippet": "API Key/CSE ID tidak valid."})
        return articles
    try:
        result = google_search_service.cse().list(q=query, cx=GOOGLE_CSE_ID, num=3, gl='id').execute() 
        if 'items' in result:
            for item in result['items']: articles.append({"source": item.get('displayLink','N/A'), "title": item.get('title','N/A'), "url": item.get('link','#'), "snippet": item.get('snippet','')})
        else: articles.append({"source": "Google", "title": "Tidak ada hasil", "url": "#", "snippet": ""})
    except Exception as e:
        app.logger.error(f"Error Google Search API: {str(e)}")
        err_msg = "Kuota Habis?" if "quota" in str(e).lower() else str(e)
        articles.append({"source": "Google", "title": "Error API", "url": "#", "snippet": err_msg})
    return articles

# --- Helper: Parsing Skor (Fokus 0-100) ---
def parse_score(text, keyword, default=0): 
    try:
        # Pola 1: Keyword [spasi opsional] [:/] [spasi opsional] [angka]
        pattern1 = rf"(?:{keyword}\s*[:\-]*\s*|\*\*\s*)(\d+)(?:\s*/\s*(?:100|7))?"
        matches1 = re.findall(pattern1, text, re.IGNORECASE | re.DOTALL)
        if matches1:
            score = int(matches1[0]) 
            app.logger.info(f"Parse '{keyword}': Pola 1 (Direct) -> {score}")
            # Asumsikan semua skor 0-100 kecuali keyword spesifik 1-7 (untuk Lentera)
            if "Skala 1-7" in keyword or (score <= 7 and "FINAL" not in keyword.upper() and "HASIL KERJA" not in keyword.upper()): 
                return max(1, min(7, score))
            return max(0, min(100, score)) 
                
        # Pola 2: [Angka] / 100 di baris keyword
        pattern2 = rf"{keyword}.*?(\d+)\s*/\s*100"
        match2 = re.search(pattern2, text, re.IGNORECASE | re.DOTALL)
        if match2: score = int(match2.group(1)); app.logger.info(f"Parse '{keyword}': Pola 2 (Fraction 100) -> {score}"); return max(0, min(100, score))
        
        # Pola 3: [Angka] / 7 di baris keyword
        pattern3 = rf"{keyword}.*?(\d)\s*/\s*7"
        match3 = re.search(pattern3, text, re.IGNORECASE | re.DOTALL)
        if match3: score_7 = int(match3.group(1)); app.logger.info(f"Parse '{keyword}': Pola 3 (Fraction 7) -> {score_7}"); return max(1, min(7, score_7))
            
        # Pola 4: Angka saja setelah keyword
        pattern4 = rf"{keyword}.*?(\d+)"
        match4 = re.search(pattern4, text, re.IGNORECASE | re.DOTALL)
        if match4:
            score = int(match4.group(1))
            app.logger.info(f"Parse '{keyword}': Pola 4 (After) -> {score}")
            if score > 7: return max(0, min(100, score))
            else: return max(1, min(7, score))

        app.logger.warning(f"Pola skor '{keyword}' tidak ditemukan. Default: {default}")
        return default
    except Exception as e:
        app.logger.error(f"Error parsing skor '{keyword}': {str(e)}. Default: {default}")
        return default
    
# --- Helper: Parsing Teks ---
def parse_recommendation(text):
    if not text: return "N/A"
    for line in text.splitlines():
        if "rekomendasi kelayakan:" in line.lower(): return line.split(":")[-1].strip().replace('*','')[:100]
    return "Tidak Ditemukan"
def parse_sentiment(text):
     if not text: return "N/A"
     for line in text.splitlines():
        if "sentimen umum:" in line.lower(): return line.split(":")[-1].strip().replace('*','')[:100]
     return "Tidak Ditemukan"

# --- Helper: Generate PDF ---
def generate_pdf_from_markdown(markdown_text, title="Profil Lentera"): # Tambah parameter title
    try:
        html_content = markdown.markdown(markdown_text, extensions=['extra', 'nl2br'])
        css = """ @page { size: A4; margin: 1.5cm; } body { font-family: sans-serif; font-size: 10pt; color: #333; } h1, h2, h3, h4 { color: #003366; } h1 { font-size: 18pt; text-align: center; margin-bottom: 20px; } h2 { font-size: 14pt; border-bottom: 1px solid #ccc; padding-bottom: 5px; margin-top: 15px; } strong { color: #00509E; } ul, ol { padding-left: 20px; } li { margin-bottom: 5px; } pre { background-color: #f0f0f0; padding: 10px; border-radius: 5px; font-family: monospace; white-space: pre-wrap; } a { color: #3b82f6; text-decoration: none; } """
        result = BytesIO()
        # Gunakan title dinamis
        pdf = pisa.CreatePDF(BytesIO(f"<html><head><style>{css}</style></head><body><h1>{title}</h1>{html_content}</body></html>".encode('utf-8')), dest=result)
        if pdf.err: app.logger.error(f"Error xhtml2pdf: {pdf.err}"); return None
        result.seek(0); return result
    except Exception as e: app.logger.error(f"Error generate_pdf: {str(e)}"); return None

# --- MODUL 1: LENTERA ---
@app.route('/api/lentera/generate-case', methods=['POST'])
def lentera_generate_case():
    app.logger.info(f"LENTERA Generate Case - Origin: {request.origin}")
    provider = request.form.get('provider', 'gemini')
    jabatan = request.form.get('jabatan', '') 
    cv_text = "" 
    filename = "N/A"

    if not jabatan: 
        app.logger.error("LENTERA Generate Case Gagal: Jabatan kosong.")
        return jsonify({"error": "Jabatan yang dituju wajib diisi."}), 400

    # Proses CV jika ada (Logika sama seperti sebelumnya)
    if 'file_cv' in request.files:
        file_cv = request.files['file_cv']
        if file_cv and file_cv.filename != '': 
            filename = secure_filename(file_cv.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            try:
                file_cv.save(file_path); cv_text_extracted = extract_text_from_file(file_path); os.remove(file_path)
                if cv_text_extracted: cv_text = cv_text_extracted; app.logger.info(f"CV {filename} OK.")
                else: app.logger.warning(f"Gagal ekstrak CV {filename}.")
            except Exception as e: app.logger.error(f"Error proses CV {filename}: {str(e)}")
                 
    app.logger.info(f"LENTERA (Tahap 1): Generate Case. Jabatan: {jabatan}, CV: {'Ada' if cv_text else 'Tidak Ada'}")
    
    # --- PERBAIKAN PROMPT V3.4 ---
    
    # Persiapan bagian prompt berdasarkan ada/tidaknya CV
    if cv_text:
        cv_info_prompt = f"\n- Ringkasan CV Kandidat: {cv_text[:1500]}"
        tugas_konteks = f"berdasarkan **jabatan yang dituju** DAN **profil CV kandidat** berikut."
    else:
        cv_info_prompt = "\n- CV Kandidat: Tidak disediakan."
        tugas_konteks = f"berdasarkan **jabatan spesifik yang dituju** berikut." # Tekankan jabatan jika CV kosong

    # Prompt Baru yang Lebih Tegas
    prompt = f"""
    Anda adalah Asesor AI BKN yang sangat ahli dalam membuat soal studi kasus **spesifik** untuk jabatan ASN.

    DATA KANDIDAT:
    - Jabatan yang **HARUS** menjadi fokus utama soal: **{jabatan}** {cv_info_prompt}

    TUGAS UTAMA:
    Buat 1 (satu) soal studi kasus yang **SANGAT SPESIFIK dan RELEVAN** {tugas_konteks}
    Soal ini WAJIB menguji **nalar, problem-solving, dan integritas** yang paling krusial untuk jabatan **{jabatan}**. 
    JANGAN membuat soal generik. Fokus pada tantangan nyata yang mungkin dihadapi oleh seorang **{jabatan}**.

    Format Output (HANYA soalnya, TANPA analisis/profil):
    **Skenario:**
    [Skenario detail dan kontekstual yang Anda buat untuk jabatan '{jabatan}', sekitar 2-4 paragraf]
    
    **Pertanyaan (3 poin spesifik untuk '{jabatan}'):**
    1. [Pertanyaan relevan 1 yang menguji aspek kunci jabatan]
    2. [Pertanyaan relevan 2 yang berbeda]
    3. [Pertanyaan relevan 3, mungkin terkait dilema etis spesifik peran '{jabatan}']
    """
    # --------------------------------------------------------

    # Tambahkan Log untuk Debugging Prompt
    app.logger.info(f"Prompt AI Generate Case:\n{prompt[:500]}...") # Log 500 karakter pertama prompt

    # Panggil AI
    if provider == 'gemini': 
        result = call_gemini_api(prompt)
    elif provider == 'byteplus' and BYTEPLUS_KEY:
        result = call_byteplus_api(prompt, "Anda adalah Asesor AI BKN pembuat soal studi kasus spesifik.")
    else: 
        app.logger.error(f"Provider AI '{provider}' tidak valid/dikonfigurasi.")
        result = f"Error: Provider AI '{provider}' tidak valid."

    app.logger.info("LENTERA Generate Case - Selesai.")
    return jsonify({"case_study": result, "cv_text_cache": cv_text})

#Sub Modul Lentera - Penilaian
@app.route('/api/lentera/grade-final', methods=['POST'])
def lentera_grade_final():
    app.logger.info(f"LENTERA Grade Final - Request Origin: {request.origin}, Content-Type: {request.content_type}")
    
    # Ambil data teks dari form DULU
    provider = request.form.get('provider', 'gemini')
    jabatan = request.form.get('jabatan', 'Analis Kebijakan')
    cv_text = request.form.get('cv_text_cache', '') 
    case_study = request.form.get('case_study', '') 
    # Ambil jawaban teks dari form (mungkin ada meskipun file diupload)
    answer_text_from_form = request.form.get('answer_text', '') 
    
    answer_text = "" # Inisialisasi teks jawaban final
    filename_ans = "N/A"

    # --- LOGIKA PEMBACAAN JAWABAN (REVISI) ---
    # Prioritaskan file jika ada DAN berhasil dibaca
    if 'file_answer' in request.files:
        file_answer = request.files['file_answer']
        if file_answer and file_answer.filename != '':
            filename_ans = secure_filename(file_answer.filename)
            filepath_ans = os.path.join(app.config['UPLOAD_FOLDER'], filename_ans)
            app.logger.info(f"Mencoba membaca file jawaban: {filename_ans}")
            try:
                file_answer.save(filepath_ans)
                answer_text_extracted = extract_text_from_file(filepath_ans)
                os.remove(filepath_ans)
                if answer_text_extracted: 
                    answer_text = answer_text_extracted # Gunakan teks dari file
                    app.logger.info(f"Jawaban berhasil diekstrak dari file {filename_ans}.")
                else: 
                    app.logger.warning(f"File jawaban {filename_ans} kosong atau gagal diekstrak.")
            except Exception as e: 
                app.logger.error(f"Error saat memproses file jawaban {filename_ans}: {str(e)}")
                # Jangan return error dulu, coba fallback ke teks
    
    # Jika teks jawaban dari file KOSONG (karena tidak ada file ATAU file gagal dibaca),
    # GUNAKAN teks dari form field 'answer_text'
    if not answer_text:
        app.logger.info("Tidak ada teks dari file jawaban, menggunakan teks dari form field.")
        answer_text = answer_text_from_form # Gunakan teks dari form
        filename_ans = "Input Teks" # Update nama sumber

    # FINAL CHECK: Jika answer_text masih kosong setelah semua usaha, baru return 400
    if not answer_text:
        app.logger.error("Gagal mendapatkan teks jawaban dari file maupun form.")
        return jsonify({"error": "Tidak ada jawaban yang valid terdeteksi (baik file maupun teks)."}), 400
    # --------------------------------------------------
    app.logger.info(f"LENTERA (Tahap 2): Melanjutkan Grade Final. Sumber Jawaban: {filename_ans}")

    # --- Sisa logika (OSINT, Prompt AI, Parsing, Simpan Log) TETAP SAMA ---
    nama_kandidat = "Kandidat" 
    match_nama = re.search(r"(?:nama|name)\s*[:\-]*\s*(.+)", cv_text, re.IGNORECASE) if cv_text else None
    if match_nama: nama_kandidat = match_nama.group(1).strip().splitlines()[0] 
        
    osint_results = run_osint_analysis(f'"{nama_kandidat}" ASN OR PNS OR BKN')
    osint_summary = "\n".join([f"- [{res['source']}]: {res['snippet'][:100]}..." for res in osint_results])

    # Pastikan 'answer_text' (yang sudah final) digunakan di prompt
    prompt = f"""Anda AI Grader BKN modern...\n--- DATA ---\nJabatan:{jabatan}\nNama:{nama_kandidat}\nCV:{'(Ada)' if cv_text else '(Tidak)'} {cv_text[:1000] if cv_text else ''}\nSoal:{case_study}\nJawaban:{answer_text}\nOSINT:{osint_summary}\n--- END DATA ---\nINSTRUKSI: Buat Profil Potensi LENTERA (Markdown, Skor 1-7)...\n---...\n## 👤 PROFIL POTENSI LENTERA\n...\n### 📊 SKOR ATRIBUT (Skala 1-7):\n...\n### 📈 REKOMENDASI & PENGEMBANGAN:\n..."""
    if provider == 'gemini': result_text = call_gemini_api(prompt)
    else: result_text = call_byteplus_api(prompt, "Anda AI Grader BKN pembuat profil modern.")
    
    # Parsing Skor & Rekomendasi (Sama seperti V3)
    skor_potensi_final = parse_score(result_text, "SKOR TOTAL POTENSI")
    rekomendasi_final = parse_recommendation(result_text)
    scores_structured_dict = {
        "kualifikasi": parse_score(result_text, "Kualifikasi & Pengetahuan", default=4),
        "nalar": parse_score(result_text, "Nalar & Logika", default=4),
        "problem": parse_score(result_text, "Problem Solving", default=4),
        "osint": parse_score(result_text, "Jejak Digital", default=4),
        "integritas": parse_score(result_text, "Potensi Integritas", default=4)
    }
    # Simpan Log (Sama seperti V3)
    try: 
        log_entry = AssessmentLog(
            module='Lentera', 
            candidate_name=nama_kandidat, 
            jabatan_or_program=jabatan, 
            skor_potensi=skor_potensi_final, 
            recommendation=rekomendasi_final
        )
        db.session.add(log_entry) # Baris baru
        db.session.commit() # Baris baru
        app.logger.info(f"Log Lentera disimpan. ID: {log_entry.id}")
    except Exception as e: 
        db.session.rollback()
        app.logger.error(f"Gagal log Lentera: {str(e)}")        
    return jsonify({ "grading_result": result_text, "skor_potensi": skor_potensi_final, "scores_structured": scores_structured_dict, "recommendation": rekomendasi_final })

#Sub Modul Lentera - Eksport Penilaian
@app.route('/api/lentera/export-pdf', methods=['POST'])
def export_lentera_pdf(): 
     data = request.json; markdown_profile = data.get('profile_markdown', ''); nama_kandidat = data.get('nama_kandidat', 'Kandidat')
     if not markdown_profile: return jsonify({"error": "Data profil tidak ada"}), 400
     app.logger.info(f"LENTERA: Generate PDF untuk {nama_kandidat}")
     pdf_buffer = generate_pdf_from_markdown(markdown_profile, title="Profil Potensi Lentera") # Kirim title
     if pdf_buffer:
        safe_filename = secure_filename(f"Profil_Lentera_{nama_kandidat}.pdf")
        return send_file(pdf_buffer, mimetype='application/pdf', as_attachment=True, download_name=safe_filename)
     else: return jsonify({"error": "Gagal generate PDF"}), 500

# --- MODUL 2: SELAYAR ---
@app.route('/api/selayar/osint-sentiment', methods=['POST'])
def selayar_osint_sentiment():
    # --- PERBAIKAN NameError V3.8 ---
    data = request.json; program_kerja, provider = data.get('program', 'Pelayanan Publik'), data.get('provider', 'byteplus')
    app.logger.info(f"SELAYAR (OSINT): Dipicu. Program: {program_kerja}, Provider: {provider}")
    osint_articles = run_osint_analysis(program_kerja) 
    
    sentiment_summary = "Error: Gagal memproses sentimen."
    skor_sentimen = 0
    sentimen_umum = "Error" # <-- Inisialisasi default
    
    if not osint_articles or any("Error" in article.get("title", "") for article in osint_articles):
        app.logger.error("SELAYAR OSINT Gagal mendapatkan artikel.")
        sentiment_summary = "Error: Gagal mendapatkan data berita dari Google Search."
    else:
        osint_snippets = "\n".join([f"- [{a['source']}]: {a['snippet'][:150]}..." for a in osint_articles if a.get('snippet')])
        prompt = f"""Anda AI Analis Sentimen Publik...\nOSINT Snippets "{program_kerja}":\n---\n{osint_snippets or "N/A"}\n---\nTUGAS: Format Markdown:\n**Analisis Sentimen Publik (OSINT):**\n* **Sentimen Umum:** [Positif/Negatif/Netral/Campuran]\n* **Skor Sentimen (Estimasi):** [Angka 1-100] / 100\n* **Ringkasan Utama:** [1 kalimat]"""
        
        if provider == 'gemini': sentiment_summary_raw = call_gemini_api(prompt)
        else: sentiment_summary_raw = call_byteplus_api(prompt, "Anda analis sentimen publik ringkas.")
        
        if "Error:" not in sentiment_summary_raw:
             sentiment_summary = sentiment_summary_raw
             skor_sentimen = parse_score(sentiment_summary, "Skor Sentimen", default=50) 
             sentimen_umum = parse_sentiment(sentiment_summary) # <-- Definisi di sini
        else:
             sentiment_summary = sentiment_summary_raw
             sentimen_umum = "Error (AI Call Failed)" # <-- Beri nilai jika AI error
    
    # Simpan Log DB
    try:
        log_entry = AssessmentLog(module='Selayar-OSINT', jabatan_or_program=program_kerja, sentiment=sentimen_umum) 
        db.session.add(log_entry); db.session.commit()
    except Exception as e: db.session.rollback(); app.logger.error(f"Gagal log Selayar-OSINT: {str(e)}")
    
    app.logger.info(f"SELAYAR OSINT Selesai. Sentimen: {sentimen_umum}, Skor: {skor_sentimen}") 
    return jsonify({"program": program_kerja, "sentiment_analysis_text": sentiment_summary, "sentiment_score": skor_sentimen, "articles": osint_articles})

# --- FUNGSI ANALISIS SKP (PROMPT BARU V3.8) ---
@app.route('/api/selayar/analyze-skp', methods=['POST']) 
def selayar_analyze_skp():
    if 'file_skp' not in request.files: return jsonify({"error": "File SKP tidak ada"}), 400
    file_skp, provider = request.files['file_skp'], request.form.get('provider', 'gemini')
    filename = secure_filename(file_skp.filename); file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename); file_skp.save(file_path)
    app.logger.info(f"SELAYAR (SKP V3.8): Analyze Artifact. File: {filename}")
    doc_text_full = extract_text_from_file(file_path); os.remove(file_path)
    if not doc_text_full: return jsonify({"error": "Gagal baca teks SKP"}), 500
    doc_text = doc_text_full[:8000] 
    if len(doc_text_full) > 8000: app.logger.warning(f"SKP {filename} dipotong.")

    # --- PROMPT BARU V3.8 (Revisi berdasarkan umpan balik) ---
    prompt = f"""
    Anda adalah **Asesor Kinerja ASN (PPK) berbasis AI** yang **sangat teliti, objektif, dan analitis**.
    Tugas Anda adalah menganalisis dokumen SKP/EKP berikut (yang berisi Rencana, Target, Realisasi, dan Umpan Balik Pimpinan) dan menghasilkan **Analisis Objektif Kinerja SELAYAR** dalam format Markdown terstruktur.

    --- DOKUMEN SKP/EKP ---
    {doc_text}
    --- AKHIR DOKUMEN ---

    INSTRUKSI ANALISIS & PEMBUATAN PROFIL (WAJIB DIIKUTI SECARA PERSIS):
    1.  **Analisis Hasil Kerja (WAJIB DETAIL):**
        * Identifikasi **minimal 3 Rencana Hasil Kerja (RHK)** utama.
        * Untuk setiap RHK, **bandingkan** data **TARGET** (cth: "12 Dokumen") dengan **REALISASI** (cth: "3 berdasarkan Arsip"). Hitung persentase pencapaian jika mungkin (cth: 3/12 = 25%).
        * Tulis hasil analisis per RHK dalam format list, sertakan perbandingan Target vs Realisasi.
        * Berikan **Estimasi Skor Pencapaian Hasil Kerja (AI):** **WAJIB** format `[Angka 0-100] / 100`. Skor ini harus mencerminkan persentase pencapaian target DI ATAS EKSPEKTASI (Realisasi > Target, cth: 4590 vs 2000 = 100+), SESUAI EKSPEKTASI (Realisasi >= Target), atau DI BAWAH EKSPEKTASI (Realisasi < Target, cth: 3 vs 12).
    2.  **Analisis Perilaku Kerja (WAJIB):**
        * Untuk **setiap 7 aspek** (Berorientasi Pelayanan, Akuntabel, Kompeten, Harmonis, Loyal, Adaptif, Kolaboratif), baca **Umpan Balik Pimpinan** (cth: "selalu dapat diandalkan", "inovasi dan kreativitas tinggi").
        * Berdasarkan umpan balik kualitatif tersebut, berikan **Skor Perilaku** (skala 1-100). Gunakan panduan ini: DI ATAS EKSPEKTASI (cth: "inovasi tinggi", "role model") = 90-100, SESUAI EKSPEKTASI (cth: "sesuai target", "baik") = 76-89, DI BAWAH EKSPEKTASI (cth: "perlu perbaikan", tidak konsisten) = <76.
        * **WAJIB** gunakan format `* **[Aspek]:** [Angka 1-100] / 100`.
        * Hitung **Rata-rata Skor Perilaku**. **WAJIB** format `[Angka 0-100] / 100`.
    3.  **Hitung Skor Akhir (WAJIB):**
        * Hitung **SKOR KINERJA FINAL** = (60% * Skor Hasil Kerja) + (40% * Rata-rata Skor Perilaku). Bulatkan. **WAJIB** format `**[Angka 0-100] / 100**`.
    4.  **Rekomendasikan Predikat Objektif (WAJIB):**
        * Berdasarkan SKOR KINERJA FINAL, tentukan **Predikat Objektif** (Sangat Baik >90, Baik 76-90, Butuh Perbaikan 61-75, Kurang 51-60, Sangat Kurang <=50). **WAJIB** format `**[Predikat]**`.
    5.  **Bandingkan Predikat:** Sebutkan Predikat Kinerja yang **tertulis di dokumen** (cari 'PREDIKAT KINERJA PEGAWAI') dan bandingkan dengan rekomendasi objektif Anda (Sesuai/Tidak Sesuai).
    6.  **Catatan & Saran (WAJIB SELALU ADA):** Berikan **Catatan Asesor AI** dan **Saran Pengembangan** (actionable), **meskipun** kinerjanya sudah Baik/Sangat Baik.

    --- TEMPLATE OUTPUT MARKDOWN (HARUS DIIKUTI) ---
    ## 📋 ANALISIS OBJEKTIF KINERJA SELAYAR

    **Pegawai:** [Deteksi Nama Pegawai, cth: ANDREAS GHANNESON NAINGGOLAN, S.Kom]
    **Periode:** [Deteksi Periode Penilaian, cth: 1 OKTOBER SD 31 DESEMBER TAHUN 2023]
    **Dokumen:** {filename}

    ---
    ### 🎯 ANALISIS HASIL KERJA (TARGET vs REALISASI):
    * **RHK 1 (cth: Perekapan Data Anjab):** Target: [Target, cth: 2000 Data]. Realisasi: [Realisasi, cth: 4590 Data] (Pencapaian: **Di Atas Ekspektasi**). Umpan Balik: [Feedback, cth: "Sangat Baik Sekali"].
    * **RHK 2 (cth: Penyusunan Hasil Anjab):** Target: [Target, cth: 2000 Data]. Realisasi: [Realisasi, cth: 4590 Jabatan] (Pencapaian: **Di Atas Ekspektasi**). Umpan Balik: [Feedback, cth: "Pertahankan"].
    * **RHK 3 (cth: Kebutuhan Jabatan):** Target: [Target, cth: 272 Kebutuhan]. Realisasi: [Realisasi, cth: 47 Data Jabatan] (Pencapaian: **Di Bawah Ekspektasi**). Umpan Balik: [Feedback, cth: "Kerja Bagus"].
    * **Ringkasan Umum Hasil Kerja:** [Kesimpulan singkat, cth: Sebagian besar hasil kerja utama melampaui target, namun ada 1 RHK yang realisasinya jauh di bawah target.]
    * **Estimasi Skor Pencapaian Hasil Kerja (AI):** [Angka 0-100] / 100

    ---
    ### ✨ ANALISIS PERILAKU KERJA (Skala 0-100):
    * **Berorientasi Pelayanan:** [Angka 1-100] / 100
    * **Akuntabel:** [Angka 1-100] / 100
    * **Kompeten:** [Angka 1-100] / 100
    * **Harmonis:** [Angka 1-100] / 100
    * **Loyal:** [Angka 1-100] / 100
    * **Adaptif:** [Angka 1-100] / 100
    * **Kolaboratif:** [Angka 1-100] / 100
    * **Rata-rata Skor Perilaku:** [Hitung rata-rata] / 100
    * **Ringkasan Umpan Balik Pimpinan (Perilaku):** [Ringkasan singkat, cth: Pimpinan memberikan umpan balik sangat positif, "kualitas terbaik", "inovasi tinggi".]

    ---
    ### 💯 SKOR KINERJA FINAL (Estimasi AI):
    *(Bobot: 60% Hasil Kerja + 40% Rata-rata Perilaku)*
    **[Angka 0-100] / 100**
    ---
    ### 🏆 REKOMENDASI PREDIKAT KINERJA (Objektif):
    **[Predikat: Sangat Baik / Baik / Butuh Perbaikan / Kurang / Sangat Kurang]**
    * *Justifikasi AI:* [1 kalimat alasan berdasarkan skor final, cth: Skor final tinggi didukung oleh hasil kerja di atas ekspektasi dan perilaku yang sangat positif.]
    * **Perbandingan dengan Dokumen:** Predikat di dokumen: '[Predikat dari Dokumen, cth: BAIK]'. Rekomendasi AI: [Sesuai/Tidak Sesuai].

    ---
    ### 💡 CATATAN & SARAN PENGEMBANGAN (AI):
    * **Catatan Asesor AI:** [Sebutkan kekuatan, cth: Menunjukkan inisiatif tinggi dan hasil kerja melebihi target di beberapa area kunci.]
    * **Saran Pengembangan:** [Berikan saran, cth: Pertahankan kinerja ekselen dan fokus pada peningkatan RHK Kebutuhan Jabatan agar sesuai target di periode berikutnya.]

    ---
    *Disclaimer: Analisis ini adalah estimasi AI.*
    ---
    """
    # --------------------------------------------------------
    
    if provider == 'gemini': result_text = call_gemini_api(prompt)
    else: result_text = call_byteplus_api(prompt, "Anda Asesor Kinerja ASN Objektif dan Analitis.")
    
    if "Error:" in result_text: 
        app.logger.error(f"Panggilan AI gagal SKP {filename}: {result_text}")
        return jsonify({"artifact_analysis": result_text, "skor_kinerja": 0, "scores_structured": {}}), 500
    
    skor_kinerja_final = parse_score(result_text, "SKOR KINERJA FINAL", default=0) 
    scores_perilaku = {
        "pelayanan": parse_score(result_text, "Berorientasi Pelayanan", default=0),
        "akuntabel": parse_score(result_text, "Akuntabel", default=0),
        "kompeten": parse_score(result_text, "Kompeten", default=0),
        "harmonis": parse_score(result_text, "Harmonis", default=0),
        "loyal": parse_score(result_text, "Loyal", default=0),
        "adaptif": parse_score(result_text, "Adaptif", default=0),
        "kolaboratif": parse_score(result_text, "Kolaboratif", default=0)
    }
    scores_structured_dict = scores_perilaku 
    nama_pegawai_skp = filename.split('.')[0]
    try:
        log_entry = AssessmentLog(module='Selayar-SKP', candidate_name=nama_pegawai_skp, jabatan_or_program="Analisis SKP", skor_kinerja=skor_kinerja_final)
        db.session.add(log_entry); db.session.commit()
    except Exception as e: db.session.rollback(); app.logger.error(f"Gagal log Selayar-SKP: {str(e)}")
    return jsonify({ "artifact_analysis": result_text, "skor_kinerja": skor_kinerja_final, "scores_structured": scores_structured_dict })

@app.route('/api/selayar/export-pdf', methods=['POST'])
def export_selayar_pdf():
    data = request.json; markdown_profile = data.get('profile_markdown', ''); nama_file_skp = data.get('nama_file_skp', 'SKP_Pegawai')
    if not markdown_profile: return jsonify({"error": "Data profil SKP tidak ada"}), 400
    app.logger.info(f"SELAYAR: Generate PDF untuk {nama_file_skp}")
    pdf_buffer = generate_pdf_from_markdown(markdown_profile, title="Analisis Objektif Kinerja SELAYAR") 
    if pdf_buffer:
        safe_filename = secure_filename(f"Analisis_Selayar_{nama_file_skp.replace('.pdf','').replace('.txt','')}.pdf")
        return send_file(pdf_buffer, mimetype='application/pdf', as_attachment=True, download_name=safe_filename)
    else: return jsonify({"error": "Gagal generate PDF SKP"}), 500


@app.route('/api/history', methods=['GET'])
def get_history():
    """Mengambil semua log asesmen dari database."""
    try:
        logs = AssessmentLog.query.order_by(AssessmentLog.timestamp.desc()).all()
        history_data = [{
            "id": log.id,
            "timestamp": log.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
            "module": log.module,
            "candidate_name": log.candidate_name,
            "jabatan_or_program": log.jabatan_or_program,
            "skor_potensi": log.skor_potensi,
            "skor_kinerja": log.skor_kinerja,
            "recommendation": log.recommendation,
            "sentiment": log.sentiment
        } for log in logs]
        return jsonify(history_data)
    except Exception as e:
        app.logger.error(f"Gagal mengambil history: {str(e)}")
        return jsonify({"error": "Gagal mengambil riwayat asesmen"}), 500

@app.route('/api/history/<log_id>', methods=['DELETE'])
def delete_history_entry(log_id):
    """Menghapus entri log berdasarkan ID."""
    try:
        log_entry = AssessmentLog.query.get(log_id)
        if log_entry:
            db.session.delete(log_entry)
            db.session.commit()
            app.logger.info(f"Log entry {log_id} dihapus.")
            return jsonify({"success": True, "message": "Entri riwayat dihapus."})
        else:
            return jsonify({"success": False, "message": "Entri tidak ditemukan."}), 404
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Gagal menghapus log entry {log_id}: {str(e)}")
        return jsonify({"error": "Gagal menghapus entri riwayat"}), 500
# -----------------------------

# --- MODUL 3: NAKHODA (Sama seperti V2.1) ---
def analyze_graph(pegawai_list, kolaborasi_list):
    G = nx.Graph()
    for p in pegawai_list:
        potensi = p.get('skor_potensi', 50); kinerja = p.get('skor_kinerja', 50)
        combined_score = (kinerja * 0.6) + (potensi * 0.4)
        G.add_node(p['id'], label=p.get('nama','N/A'), unit=p.get('unit','N/A'), jabatan=p.get('jabatan','N/A'), score=combined_score)
    for k in kolaborasi_list:
        if G.has_node(k.get('source')) and G.has_node(k.get('target')):
            G.add_edge(k['source'], k['target'], label=k.get('project',''))
    if not G.nodes: return {"nodes": [], "edges": [], "metrics": {"total_pegawai": 0, "total_kolaborasi": 0, "avg_effectiveness": 0, "num_silos": 0}}
    centrality = nx.degree_centrality(G)
    num_silos = nx.number_connected_components(G)
    total_effectiveness_score = sum((G.nodes[n].get('score', 50) * (1 + centrality.get(n, 0))) for n in G.nodes())
    nodes_for_reactflow = []
    for node in G.nodes(data=True):
        node_id, node_data = node; node_score = node_data.get('score', 50)
        bg = '#90EE90' if node_score > 80 else ('#FFD700' if node_score > 60 else '#F08080')
        nodes_for_reactflow.append({"id": node_id, "position": {"x": 0, "y": 0},"data": {"label": f"{node_data.get('label','N/A')} ({node_data.get('unit','N/A')})\nSkor: {node_score:.0f}"},"style": { "background": bg, "border": "1px solid #333", "whiteSpace": "pre-line", "textAlign": "center"}})
    edges_for_reactflow = [{"id": f"e-{e[0]}-{e[1]}", "source": e[0], "target": e[1], "label": e[2].get('label', ''), "animated": True} for e in G.edges(data=True)]
    metrics = {"total_pegawai": G.number_of_nodes(), "total_kolaborasi": G.number_of_edges(), "avg_effectiveness": (total_effectiveness_score / G.number_of_nodes()) if G.number_of_nodes() > 0 else 0, "num_silos": num_silos}
    return {"nodes": nodes_for_reactflow, "edges": edges_for_reactflow, "metrics": metrics}


@app.route('/api/nakhoda/get-graph', methods=['GET'])
def get_graph():
    app.logger.info("Modul NAKHODA: Mengambil graf awal (dummy_data_v2.json).")
    try:
        with open('dummy_data_v2.json', 'r') as f: data = json.load(f)
    except FileNotFoundError:
        app.logger.warning("dummy_data_v2.json tidak ada! Membuat baru...")
        new_dummy_data={"pegawai": [{"id": "1", "nama": "Anya", "unit": "A", "skor_potensi": 90, "skor_kinerja": 95}, {"id": "2", "nama": "Budi", "unit": "A", "skor_potensi": 95, "skor_kinerja": 70}, {"id": "3", "nama": "Citra", "unit": "SDM", "skor_potensi": 80, "skor_kinerja": 85}], "kolaborasi": [{"source": "1", "target": "2"}, {"source": "1", "target": "3"}, {"source": "2", "target": "3"}]}
        with open('dummy_data_v2.json', 'w') as f: json.dump(new_dummy_data, f, indent=2)
        data = new_dummy_data
    result = analyze_graph(data['pegawai'], data['kolaborasi'])
    return jsonify(result)

@app.route('/api/nakhoda/load-custom-graph', methods=['POST'])
def load_custom_graph():
    data = request.json
    pegawai_json_text, kolaborasi_json_text = data.get('pegawaiData'), data.get('kolaborasiData')
    app.logger.info("Modul NAKHODA: Menerima data graf kustom.")
    try:
        pegawai_list, kolaborasi_list = json.loads(pegawai_json_text), json.loads(kolaborasi_json_text)
        result = analyze_graph(pegawai_list, kolaborasi_list)
        return jsonify(result)
    except Exception as e: return jsonify({"error": str(e)}), 400

@app.route('/api/nakhoda/simulate-move', methods=['POST'])
def simulate_move():
    try:
        sim_data = request.json
        pegawai_id, target_unit = sim_data.get('pegawaiId'), sim_data.get('targetUnit')
        pegawai_list, kolaborasi_list = sim_data.get('pegawaiList', []), sim_data.get('kolaborasiList', [])
        app.logger.info(f"Modul NAKHODA: Simulasi pemindahan {pegawai_id} ke {target_unit}")
        if not pegawai_id or not target_unit: return jsonify({"error": "Data simulasi tidak lengkap"}), 400
        if not pegawai_list:
             with open('dummy_data_v2.json', 'r') as f: data = json.load(f)
             pegawai_list, kolaborasi_list = data['pegawai'], data['kolaborasi']
        sim_pegawai_list = [p.copy() if p['id'] != pegawai_id else {**p, 'unit': target_unit} for p in pegawai_list]
        sim_result = analyze_graph(sim_pegawai_list, kolaborasi_list)
        original_result = analyze_graph(pegawai_list, kolaborasi_list)
        original_score, new_score = original_result['metrics']['avg_effectiveness'], sim_result['metrics']['avg_effectiveness']
        impact = new_score - original_score
        impact_report = (f"Laporan Dampak Simulasi:\n- Pegawai (ID: {pegawai_id}) dipindah ke Unit '{target_unit}'.\n\n**Skor Efektivitas Tim (Baru): {new_score:.2f}**\n**Skor Efektivitas Tim (Lama): {original_score:.2f}**\n**Dampak Perubahan: {impact:+.2f} Poin**\n\n- Jumlah Silo Organisasi: {sim_result['metrics']['num_silos']} (Sebelum: {original_result['metrics']['num_silos']}")
        sim_result['report'] = impact_report
        app.logger.info("Simulasi berhasil.")
        return jsonify(sim_result)
    except Exception as e: return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5001)